"""Report rendering and export utilities."""

import io
import io
import re
import html as _html
import unicodedata

import streamlit as st

from utils.helpers import md_to_html, fix_mojibake

# ── Export helpers ────────────────────────────────────────────────────────────

def report_to_docx_bytes(report_md: str, title: str = "AnalystAI Report") -> bytes:
    """Convert markdown report to a Word .docx file and return bytes."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Normalize encoding and strip leading decorative symbols from headings
        try:
            report_md = fix_mojibake(report_md)
        except Exception:
            pass

        def _strip_leading_symbols(s: str) -> str:
            s = s.lstrip()
            # Remove leading symbol characters (emojis/icons) while they appear
            while s:
                cat = unicodedata.category(s[0])
                if cat.startswith("So") or cat.startswith("Sk"):
                    s = s[1:]
                    continue
                break
            return s.lstrip()

        # Title
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Parse markdown sections
        # Unescape any HTML entities (e.g., &nbsp;, &amp;) that may have crept in
        try:
            report_md = _html.unescape(report_md)
        except Exception:
            pass

        lines = report_md.splitlines()
        for line in lines:
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue
            if stripped.startswith("## "):
                doc.add_heading(_strip_leading_symbols(stripped[3:].strip()), level=1)
            elif stripped.startswith("### "):
                doc.add_heading(_strip_leading_symbols(stripped[4:].strip()), level=2)
            elif stripped.startswith("# "):
                doc.add_heading(_strip_leading_symbols(stripped[2:].strip()), level=1)
            elif stripped.startswith(("- ", "* ", "• ")):
                item = stripped[2:].strip()
                # Strip markdown bold
                item = re.sub(r'\*\*(.+?)\*\*', r'\1', item)
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(item)
            elif re.match(r'^\d+\.\s', stripped):
                item = re.sub(r'^\d+\.\s', '', stripped)
                item = re.sub(r'\*\*(.+?)\*\*', r'\1', item)
                p = doc.add_paragraph(style="List Number")
                p.add_run(item)
            else:
                # Regular paragraph — handle inline bold
                p = doc.add_paragraph()
                parts = re.split(r'\*\*(.+?)\*\*', stripped)
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:  # odd parts are bold
                        run.bold = True

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()
    except Exception as e:
        return b""


def report_to_pdf_bytes(report_md: str, title: str = "AnalystAI Report") -> bytes:
    """Convert markdown report to PDF bytes using HTML → PDF via weasyprint or xhtml2pdf."""
    try:
        import markdown as md_lib
        try:
            report_md = fix_mojibake(report_md)
        except Exception:
            pass
        # Sanitize common artifacts that sometimes appear when copying from
        # rich text or when fonts like Zapf Dingbats introduce glyphs.
        def _sanitize_for_pdf(md_text: str) -> str:
            if not md_text:
                return md_text
            # Remove CSS-like selectors that sometimes get prepended (e.g. "p.p1 {...}")
            md_text = re.sub(r'(?m)^[A-Za-z0-9_\.\-]+\s*\{[^}]*\}\s*$', '', md_text)
            # Remove explicit font lines referencing Zapf Dingbats or other dingbat fonts
            md_text = re.sub(r"(?i)font:\s*[^;]+Zapf\s*Dingbats[^;]*;?", "", md_text)
            # Replace black squares and other box/dingbat characters with a simple bullet or space
            md_text = md_text.replace('\u25A0', '-')
            md_text = md_text.replace('■', '-')
            # Replace multiple consecutive filler symbols (sometimes used as separators)
            md_text = re.sub(r'[\u25A0\u2022\u2023\u25AA\u25AB]{2,}', '-', md_text)
            # Normalize stray control whitespace and non-breaking spaces
            md_text = md_text.replace('\xa0', ' ')
            # Trim excessive spaces around punctuation introduced by replacements
            md_text = re.sub(r'\s{2,}', ' ', md_text)
            # Remove any leftover CSS blocks or HTML-style <style>..</style> if present
            md_text = re.sub(r'(?is)<style.*?>.*?</style>', '', md_text)
            return md_text.strip()

        try:
            report_md = _sanitize_for_pdf(report_md)
        except Exception:
            pass
        html_body = md_lib.markdown(report_md, extensions=["extra", "nl2br"])
        full_html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  body {{ font-family: Arial, sans-serif; font-size: 12pt; color: #1E293B;
          max-width: 800px; margin: 40px auto; line-height: 1.6; }}
  h1 {{ color: #0F172A; font-size: 22pt; border-bottom: 2px solid #6366F1;
        padding-bottom: 8px; margin-bottom: 16px; }}
  h2 {{ color: #1E293B; font-size: 16pt; margin-top: 24px; }}
  h3 {{ color: #334155; font-size: 13pt; }}
  ul, ol {{ padding-left: 20px; }}
  li {{ margin-bottom: 6px; }}
  strong {{ color: #0F172A; }}
  p {{ margin: 8px 0; }}
  .title {{ font-size: 26pt; font-weight: bold; color: #6366F1; margin-bottom: 4px; }}
  .subtitle {{ color: #64748B; font-size: 11pt; margin-bottom: 32px; }}
</style>
</head>
<body>
<div class="title">{title}</div>
<div class="subtitle">Generated by AnalystAI · Powered by OpenAI</div>
{html_body}
</body>
</html>"""

        # Try weasyprint first
        try:
            from weasyprint import HTML
            buf = io.BytesIO()
            HTML(string=full_html).write_pdf(buf)
            buf.seek(0)
            return buf.read()
        except Exception:
            # weasyprint may be missing or fail; continue to fallbacks
            pass

        # Try xhtml2pdf
        try:
            from xhtml2pdf import pisa
            buf = io.BytesIO()
            pisa.CreatePDF(full_html, dest=buf)
            buf.seek(0)
            return buf.read()
        except Exception:
            pass

        # Fallback: use reportlab to generate a structured PDF (headings, lists,
        # paragraphs) that better preserves the markdown layout when HTML
        # renderers are unavailable.
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
            from reportlab.lib.units import inch
            from reportlab.lib import colors

            buf = io.BytesIO()
            doc = SimpleDocTemplate(
                buf,
                pagesize=letter,
                rightMargin=40,
                leftMargin=40,
                topMargin=40,
                bottomMargin=40,
            )
            base_styles = getSampleStyleSheet()
            styles = {
                'title': ParagraphStyle('title', parent=base_styles['Title'], fontName='Helvetica-Bold', fontSize=20, textColor=colors.HexColor('#6366F1')),
                'h1': ParagraphStyle('h1', parent=base_styles['Heading1'], fontName='Helvetica-Bold', fontSize=16, textColor=colors.HexColor('#0F172A')),
                'h2': ParagraphStyle('h2', parent=base_styles['Heading2'], fontName='Helvetica-Bold', fontSize=13, textColor=colors.HexColor('#1E293B')),
                'h3': ParagraphStyle('h3', parent=base_styles['Heading3'], fontName='Helvetica-Bold', fontSize=11, textColor=colors.HexColor('#334155')),
                'body': ParagraphStyle('body', parent=base_styles['BodyText'], fontName='Helvetica', fontSize=10, leading=14, textColor=colors.HexColor('#1E293B')),
                'bullet': ParagraphStyle('bullet', parent=base_styles['BodyText'], fontName='Helvetica', fontSize=10, leftIndent=12, leading=14, textColor=colors.HexColor('#1E293B')),
            }

            story = []
            story.append(Paragraph(title, styles['title']))
            story.append(Spacer(1, 12))

            try:
                report_md = fix_mojibake(report_md)
            except Exception:
                pass

            text = _html.unescape(report_md or '')
            # Clean leftover CSS selectors and dingbats
            text = re.sub(r'(?m)^[A-Za-z0-9_\.\-]+\s*\{[^}]*\}\s*$', '', text)
            text = text.replace('■', '-')
            text = text.replace('\u25A0', '-')

            lines = text.splitlines()

            i = 0
            while i < len(lines):
                ln = lines[i].strip()
                if not ln:
                    story.append(Spacer(1, 6))
                    i += 1
                    continue

                # Headings
                if ln.startswith('# '):
                    content = ln[2:].strip()
                    content = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', content)
                    story.append(Paragraph(content, styles['h1']))
                    story.append(Spacer(1, 6))
                    i += 1
                    continue
                if ln.startswith('## '):
                    content = ln[3:].strip()
                    content = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', content)
                    story.append(Paragraph(content, styles['h2']))
                    story.append(Spacer(1, 6))
                    i += 1
                    continue
                if ln.startswith('### '):
                    content = ln[4:].strip()
                    content = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', content)
                    story.append(Paragraph(content, styles['h3']))
                    story.append(Spacer(1, 6))
                    i += 1
                    continue

                # Lists (collect consecutive list lines)
                if re.match(r'^([-*•\u2022]|\d+\.)\s+', ln):
                    items = []
                    ordered = False
                    while i < len(lines) and re.match(r'^([-*•\u2022]|\d+\.)\s+', lines[i].strip()):
                        l = lines[i].strip()
                        mnum = re.match(r'^(\d+)\.\s+(.*)', l)
                        if mnum:
                            ordered = True
                            item_text = mnum.group(2)
                        else:
                            item_text = re.sub(r'^([-*•\u2022])\s+', '', l)
                        item_text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', item_text)
                        # Remove leading dingbats
                        while item_text and unicodedata.category(item_text[0]).startswith(('So','Sk')):
                            item_text = item_text[1:]
                        items.append(Paragraph(item_text, styles['bullet']))
                        i += 1
                    lf = ListFlowable([ListItem(it) for it in items], bulletType='1' if ordered else 'bullet', start='1')
                    story.append(lf)
                    story.append(Spacer(1, 6))
                    continue

                # Regular paragraph: join subsequent non-empty, non-special lines into a paragraph
                para_lines = [ln]
                i2 = i + 1
                while i2 < len(lines) and lines[i2].strip() and not re.match(r'^(#|##|###|[-*•\u2022]|\d+\.)\s+', lines[i2].strip()):
                    para_lines.append(lines[i2].strip())
                    i2 += 1
                para = ' '.join(para_lines)
                para = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', para)
                # Strip any remaining leading symbols
                while para and unicodedata.category(para[0]).startswith(('So','Sk')):
                    para = para[1:]
                story.append(Paragraph(para, styles['body']))
                story.append(Spacer(1, 6))
                i = i2

            doc.build(story)
            buf.seek(0)
            return buf.read()
        except Exception:
            return b""
    except Exception:
        return b""


# ── Report renderer ───────────────────────────────────────────────────────────

def render_structured_report(report_md: str, df_view, kpis, cleaning_report, insights):
    """Render the markdown report with clean styling. No fragile regex parsing."""
    if not report_md:
        return '<div class="report-card"><p style="color:#64748B;">Report not yet generated.</p></div>'

    # Repair common encoding glitches before rendering
    try:
        report_md = fix_mojibake(report_md)
    except Exception:
        pass

    # Convert markdown to HTML cleanly
    try:
        import markdown as md_lib
        html_body = md_lib.markdown(report_md, extensions=["extra", "nl2br"])
    except ImportError:
        # Fallback to our own md_to_html
        html_body = md_to_html(report_md)

    # Style the HTML output
    styled = f"""
    <div style="font-family:'Inter',sans-serif;color:#1E293B;line-height:1.7;font-size:0.93rem;">
      {html_body}
    </div>
    """

    # Inject CSS to style the rendered markdown elements
    css = """
    <style>
      .report-card h1, .report-card h2 {
        color: #0F172A; font-weight: 700; margin: 1.4rem 0 0.6rem;
        padding-bottom: 6px; border-bottom: 1px solid #E2E8F0;
      }
      .report-card h3 { color: #334155; font-weight: 600; margin: 1rem 0 0.4rem; }
      .report-card p  { color: #334155; margin: 0 0 10px; line-height: 1.7; }
      .report-card ul, .report-card ol {
        padding-left: 20px; margin: 6px 0 12px;
      }
      .report-card li { color: #334155; margin-bottom: 6px; line-height: 1.6; }
      .report-card strong { color: #0F172A; }
      .report-card em { color: #475569; }
      .report-card hr { border: none; border-top: 1px solid #E2E8F0; margin: 1.2rem 0; }
    </style>
    """

    return f"{css}<div class='report-card'>{styled}</div>"
